from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches, Pt, Mm

doc = Document()  # создание экземпляра

# добавление заголовка
doc.add_heading('Отчёт за месяц', 1)
paragraph_blank = doc.add_paragraph()
paragraph = doc.add_paragraph('В этом отсчёте представлены результаты обучения:\n')
paragraph.add_run(' ключевые показатели').bold = True
# run - это нечто внутри абзаца ("параграфа"), текст, картинка, иной объект
# отступ
paragraph = doc.add_paragraph()
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# маркированный
paragraph = doc.add_paragraph('Первое', style='List Bullet')
paragraph = doc.add_paragraph('Второе', style='List Bullet')

# нумерованный
paragraph = doc.add_paragraph('Первое', style='List Number')
paragraph = doc.add_paragraph('Второе', style='List Number')

paragraph = doc.add_paragraph()
# добавим таблицу
table = doc.add_table(rows=3, cols=3)

# заполняем
for i, row in enumerate(table.rows):
    for j, cell in enumerate(table.columns):
        cell.text = f'Строка {i + 1}, Колонка {j + 1}'

doc.add_paragraph()
doc.add_picture('../images/harry_python.jpg', width=Mm(105))


doc.save('../docs/Report.docs')
